# report_agent.py
# LangGraph agent that generates a 94-page academic project report.
# Called by main.py (Cloud Function trigger).

import anthropic
import asyncio
import io
import json
import zipfile
from datetime import datetime, timedelta, timezone
from typing import TypedDict, List, Dict, Literal
from langgraph.graph import StateGraph, END
from docx import Document

client = None

def _get_client():
    global client
    if client is None:
        client = anthropic.Anthropic(max_retries=5)
    return client


# ── Helpers ───────────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_str: str):
    from docx.shared import RGBColor
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _lighten_rgb(rgb, factor: float = 0.3):
    from docx.shared import RGBColor
    r = int(rgb[0] + (255 - rgb[0]) * factor)
    g = int(rgb[1] + (255 - rgb[1]) * factor)
    b = int(rgb[2] + (255 - rgb[2]) * factor)
    return RGBColor(r, g, b)

def _add_chapter_separator(doc):
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'thick')
    top.set(qn('w:sz'),    '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'auto')
    pBdr.append(top)
    pPr.append(pBdr)


def _add_header_footer(section, job: dict, heading_rgb, accent_hex: str):
    from docx.oxml.ns   import qn
    from docx.oxml      import OxmlElement
    from docx.shared    import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section.different_first_page_header_footer = True

    # ── Header ────────────────────────────────────────────────────
    hdr  = section.header
    hdr.is_linked_to_previous = False
    para = hdr.paragraphs[0]
    para.clear()

    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)

    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), accent_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

    r_left = para.add_run(job.get('title', '')[:50])
    r_left.font.name      = 'Calibri'
    r_left.font.size      = Pt(9)
    r_left.font.bold      = True
    r_left.font.color.rgb = heading_rgb

    para.add_run('\t')

    r_right = para.add_run(job.get('student_name', '')[:40])
    r_right.font.name      = 'Calibri'
    r_right.font.size      = Pt(9)
    r_right.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Footer ────────────────────────────────────────────────────
    ftr  = section.footer
    ftr.is_linked_to_previous = False
    para = ftr.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field_run(instr: str):
        r   = para.add_run()
        beg = OxmlElement('w:fldChar')
        beg.set(qn('w:fldCharType'), 'begin')
        r._r.append(beg)
        txt = OxmlElement('w:instrText')
        txt.set(qn('xml:space'), 'preserve')
        txt.text = instr
        r._r.append(txt)
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        r._r.append(end)
        r.font.name      = 'Calibri'
        r.font.size      = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _field_run(' PAGE ')
    sep = para.add_run(' of ')
    sep.font.name      = 'Calibri'
    sep.font.size      = Pt(9)
    sep.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _field_run(' NUMPAGES ')


CHAPTER_NAMES = [
    "abstract", "literature", "requirements", "system_design",
    "implementation", "database", "testing", "results",
    "conclusion", "references"
]

CHAPTER_PROMPTS = {
    "abstract": """
Write Chapter 1: Abstract and Introduction.
Sections to cover:
- Abstract (250 words, standalone summary)
- 1.1 Background and Motivation
- 1.2 Problem Statement
- 1.3 Objectives (minimum 5 numbered points)
- 1.4 Scope of the Project
- 1.5 Significance of the Study
- 1.6 Organisation of the Report
Minimum: 900 words. Do not use markdown. Use plain paragraph text.
""",
    "literature": """
Write Chapter 2: Literature Review.
Sections to cover:
- 2.1 Introduction to the domain
- 2.2 Review of existing systems (describe 5 existing systems)
- 2.3 Comparison table (System | Key Features | Limitations)
- 2.4 Research gaps identified
- 2.5 How this project addresses the gaps
- 2.6 Summary
Minimum: 1000 words. Include a comparison table in plain text format using pipe | separators.
""",
    "requirements": """
Write Chapter 3: Requirements Analysis.
Sections to cover:
- 3.1 Introduction
- 3.2 Functional Requirements (minimum 10 numbered points)
- 3.3 Non-Functional Requirements (minimum 8 points)
- 3.4 User Roles and Access Levels
- 3.5 Use Case Descriptions (5 use cases with actor, precondition, flow)
- 3.6 System Constraints
- 3.7 Assumptions and Dependencies
Minimum: 1200 words.
""",
    "system_design": """
Write Chapter 4: System Design.
Sections to cover:
- 4.1 System Architecture Overview
- 4.2 Architecture Diagram Description (describe components and interactions)
- 4.3 Data Flow Diagram Level 0 (Context Diagram) — describe in text
- 4.4 Data Flow Diagram Level 1 — describe each process
- 4.5 Entity-Relationship Diagram Description
  (list all entities, attributes, relationships)
- 4.6 Sequence Diagram Description (main workflow)
- 4.7 Component Interaction Summary
Minimum: 1800 words.
""",
    "implementation": """
Write Chapter 5: Implementation.
Sections to cover:
- 5.1 Development Environment and Tools
- 5.2 Project Structure and File Organisation
- 5.3 Module-wise Implementation (describe each module separately)
  For each module: purpose, key functions, implementation approach
- 5.4 Key Algorithms and Logic
- 5.5 API Design and Endpoints (if applicable)
- 5.6 Security Implementation
- 5.7 Error Handling Strategy
- 5.8 Integration Between Modules
Minimum: 2500 words. This is the longest chapter.
""",
    "database": """
Write Chapter 6: Database Design.
Sections to cover:
- 6.1 Database Management System Used and Justification
- 6.2 Database Architecture
- 6.3 Entity Description (describe each entity)
- 6.4 Table Structures
  For EVERY table write the schema in this EXACT pipe-delimited format — no exceptions:
  Column | Datatype | Constraints | Description
  -------|----------|-------------|------------
  id     | INT      | PRIMARY KEY | Unique row identifier
  name   | VARCHAR(100) | NOT NULL | Full name
  (continue all columns)
  Do this for at least 4 tables.
- 6.5 Relationships and Foreign Keys
- 6.6 Normalisation (explain up to 3NF)
- 6.7 Sample Data — for 3 main tables, show 3 sample rows each in pipe-delimited format:
  column1 | column2 | column3
  --------|---------|--------
  value1  | value2  | value3
- 6.8 Indexes and Performance Considerations
IMPORTANT: Every table MUST be formatted as pipe-delimited rows as shown above.
Minimum: 1000 words.
""",
    "testing": """
Write Chapter 7: Testing.
Sections to cover:
- 7.1 Testing Strategy and Approach
- 7.2 Types of Testing Performed
- 7.3 Test Plan
- 7.4 Unit Test Cases (15 test cases) — MUST use this exact pipe-delimited format:
  Test ID | Module | Description | Input | Expected Output | Status
  --------|--------|-------------|-------|-----------------|-------
  TC-01   | Login  | Valid login | email+pass | Dashboard shown | Pass
  (continue for all 15 rows)
- 7.5 Integration Test Cases (10 test cases in same pipe-delimited format)
- 7.6 User Acceptance Testing (5 scenarios)
- 7.7 Test Results Summary
- 7.8 Defects Found and Resolved
Minimum: 1200 words.
""",
    "results": """
Write Chapter 8: Results and Discussion.
Sections to cover:
- 8.1 Introduction
- 8.2 System Screenshots Description
  (describe minimum 8 screens: login, dashboard, main features, reports)
- 8.3 Feature Demonstration
- 8.4 Performance Analysis
- 8.5 Comparison with Existing Systems
- 8.6 User Feedback Summary
- 8.7 Limitations of Current Implementation
Minimum: 1000 words.
""",
    "conclusion": """
Write Chapter 9: Conclusion and Future Work.
Sections to cover:
- 9.1 Summary of Work Done (10 key accomplishments)
- 9.2 Objectives Achieved (map back to objectives in Chapter 1)
- 9.3 Limitations of the Current System
- 9.4 Future Enhancements (minimum 7 with descriptions)
- 9.5 Learning Outcomes
- 9.6 Final Remarks
Minimum: 700 words.
""",
    "references": """
Write Chapter 10: References and Appendix.
Sections:
- 10.1 References
  Generate 15 realistic IEEE format references relevant to the
  tech stack and domain. Format:
  [1] Author, "Title," Journal/Conference, vol., no., pp., Year.
- 10.2 Glossary (15 technical terms used in the project with definitions)
- 10.3 Appendix A — Installation Guide
  (step-by-step setup instructions for the tech stack)
- 10.4 Appendix B — User Manual Summary
Minimum: 600 words.
"""
}


# ── Branch-specific chapter title overrides ──────────────────────────────────

_BRANCH_CHAPTER_TITLES = {
    "ece": {
        "requirements"   : "Chapter 3: System Specifications & Requirements",
        "system_design"  : "Chapter 4: Hardware Architecture & System Design",
        "implementation" : "Chapter 5: Hardware Implementation & Firmware",
        "database"       : "Chapter 6: Circuit Design & Simulation",
        "testing"        : "Chapter 7: Hardware Testing & Validation",
        "results"        : "Chapter 8: Experimental Results & Analysis",
    },
    "eee": {
        "requirements"   : "Chapter 3: System Specifications & Design Constraints",
        "system_design"  : "Chapter 4: Power System & Circuit Design",
        "implementation" : "Chapter 5: Hardware Construction & Programming",
        "database"       : "Chapter 6: Simulation Design & Modelling",
        "testing"        : "Chapter 7: Simulation Testing & Lab Validation",
        "results"        : "Chapter 8: Experimental Results & Performance Analysis",
    },
    "mechanical": {
        "requirements"   : "Chapter 3: Design Requirements & Specifications",
        "system_design"  : "Chapter 4: Design Methodology & CAD Modelling",
        "implementation" : "Chapter 5: Manufacturing & Fabrication",
        "database"       : "Chapter 6: Structural & Thermal Analysis",
        "testing"        : "Chapter 7: Testing & Validation",
        "results"        : "Chapter 8: Results & Performance Evaluation",
    },
    "civil": {
        "requirements"   : "Chapter 3: Site Analysis & Design Requirements",
        "system_design"  : "Chapter 4: Structural System Design",
        "implementation" : "Chapter 5: Structural Analysis & Calculations",
        "database"       : "Chapter 6: Detailed Design & Specifications",
        "testing"        : "Chapter 7: Testing & Structural Validation",
        "results"        : "Chapter 8: Design Results & Safety Assessment",
    },
}

# ── Branch-specific chapter prompt overrides ──────────────────────────────────

_BRANCH_CHAPTER_PROMPTS = {
    "ece": {
        "requirements": """
Write Chapter 3: System Specifications & Requirements.
Sections to cover:
- 3.1 Introduction
- 3.2 Functional Requirements (minimum 10 numbered points)
- 3.3 Hardware Requirements (components, sensors, microcontrollers, power supply specs)
- 3.4 Firmware / Software Requirements
- 3.5 Performance Specifications (voltage, current, frequency, accuracy targets)
- 3.6 Environmental Constraints (temperature range, humidity, operating conditions)
- 3.7 Safety and Regulatory Standards
- 3.8 Assumptions and Dependencies
Minimum: 1200 words.
""",
        "system_design": """
Write Chapter 4: Hardware Architecture & System Design.
Sections to cover:
- 4.1 System Overview and Block Diagram Description
- 4.2 Microcontroller / Processor Selection and Justification
- 4.3 Sensor and Actuator Selection (specifications, interfacing method)
- 4.4 Power Supply Design (voltage levels, current draw, regulation)
- 4.5 Communication Protocols Used (I2C, SPI, UART, Wi-Fi, Bluetooth, etc.)
- 4.6 Signal Flow Description
- 4.7 Hardware Architecture Diagram Description (describe each module)
- 4.8 System Integration Overview
Minimum: 1800 words.
""",
        "implementation": """
Write Chapter 5: Hardware Implementation & Firmware.
Sections to cover:
- 5.1 Hardware Assembly and Prototyping Steps
- 5.2 Component Pin Connections and Wiring (describe each module)
- 5.3 Firmware / Embedded Code Structure
- 5.4 Module-wise Implementation (for each hardware module: purpose, interfacing, logic)
- 5.5 Communication Protocol Implementation Details
- 5.6 Data Acquisition, Processing, and Output
- 5.7 Error Handling and Safety Mechanisms
- 5.8 System Integration and Bench Testing Setup
Minimum: 2500 words.
""",
        "database": """
Write Chapter 6: Circuit Design & Simulation.
Sections to cover:
- 6.1 Circuit Design Overview
- 6.2 Schematic Description (each block: component values, ratings, purpose)
- 6.3 Component List (Table: Component | Specification | Quantity | Purpose)
- 6.4 PCB Design Considerations (layers, trace width, grounding strategy)
- 6.5 Simulation Environment (Proteus / LTSpice / MATLAB Simulink)
- 6.6 Simulation Parameters and Setup
- 6.7 Expected vs Simulated Results Comparison
- 6.8 Circuit Optimization and Improvements Made
Minimum: 1000 words. Include component table using pipe | separators.
""",
        "testing": """
Write Chapter 7: Hardware Testing & Validation.
Sections to cover:
- 7.1 Testing Strategy and Equipment Used (oscilloscope, multimeter, logic analyser)
- 7.2 Module-level Unit Tests (Table: Module | Test | Expected | Result | Status — 15 rows)
- 7.3 Integration Testing (10 test cases in same format)
- 7.4 Sensor Calibration Procedure and Results
- 7.5 Signal Quality and Noise Analysis
- 7.6 Power Consumption Measurement
- 7.7 Issues Encountered and Resolved
- 7.8 Test Results Summary
Minimum: 1200 words.
""",
        "results": """
Write Chapter 8: Experimental Results & Analysis.
Sections to cover:
- 8.1 Introduction
- 8.2 Hardware Output Measurements (voltage, current, frequency, sensor readings)
- 8.3 Waveform Analysis (describe oscilloscope traces and signal characteristics)
- 8.4 Performance vs Design Specifications (Table: Parameter | Target | Achieved | Error %)
- 8.5 Sensor Accuracy and Response Time Analysis
- 8.6 Power Efficiency Analysis
- 8.7 Comparison with Existing Solutions
- 8.8 Limitations and Scope for Improvement
Minimum: 1000 words.
""",
    },
    "eee": {
        "requirements": """
Write Chapter 3: System Specifications & Design Constraints.
Sections to cover:
- 3.1 Introduction
- 3.2 Electrical System Requirements (voltage, current, power ratings)
- 3.3 Load Analysis and Power Demand Calculation
- 3.4 Control System Requirements
- 3.5 Protection and Safety Requirements (overcurrent, overvoltage, fault protection)
- 3.6 Regulatory and Standards Compliance (IEC, IEEE, IS standards)
- 3.7 Environmental and Operational Constraints
- 3.8 Assumptions and Dependencies
Minimum: 1200 words.
""",
        "system_design": """
Write Chapter 4: Power System & Circuit Design.
Sections to cover:
- 4.1 System Architecture Overview
- 4.2 Power Circuit Design (topology, switching elements, passive components)
- 4.3 Control System Design (feedback loops, PID / PWM control strategy)
- 4.4 Protection Circuit Design (overcurrent, overvoltage, short circuit protection)
- 4.5 Component Selection and Justification with Ratings
- 4.6 Single Line Diagram Description
- 4.7 Control Signal Flow and Block Diagram
- 4.8 MATLAB Simulink Model Overview
Minimum: 1800 words.
""",
        "implementation": """
Write Chapter 5: Hardware Construction & Programming.
Sections to cover:
- 5.1 Hardware Construction and Panel Wiring Steps
- 5.2 PCB or Panel Board Layout Description
- 5.3 PLC / Microcontroller Programming (describe ladder logic or code structure)
- 5.4 SCADA / HMI Configuration (if applicable)
- 5.5 MATLAB Simulink Model Implementation Details
- 5.6 Control Algorithm Implementation
- 5.7 System Commissioning and Integration Steps
- 5.8 Safety Procedures Followed During Implementation
Minimum: 2500 words.
""",
        "database": """
Write Chapter 6: Simulation Design & Modelling.
Sections to cover:
- 6.1 Simulation Tool Selection and Justification
- 6.2 MATLAB Simulink / PSIM / PSCAD Model Description
- 6.3 Simulation Parameters (Table: Parameter | Symbol | Value | Unit)
- 6.4 Transfer Function and Mathematical Modelling
- 6.5 Simulation Scenarios and Test Cases Defined
- 6.6 Controller Design and Tuning (PID gains, switching frequency)
- 6.7 Expected Simulation Results
- 6.8 Model Validation Approach
Minimum: 1000 words. Include parameters table using pipe | separators.
""",
        "testing": """
Write Chapter 7: Simulation Testing & Lab Validation.
Sections to cover:
- 7.1 Testing Methodology (simulation + hardware)
- 7.2 Simulation Test Cases (Table: Parameter | Input | Expected | Simulated | Status — 15 rows)
- 7.3 Hardware Lab Test Cases (10 cases in same format)
- 7.4 Steady-State Performance Analysis
- 7.5 Transient Response Analysis (step response, load change)
- 7.6 Fault Condition Testing
- 7.7 Simulation vs Hardware Comparison
- 7.8 Issues Found and Resolved
Minimum: 1200 words.
""",
        "results": """
Write Chapter 8: Experimental Results & Performance Analysis.
Sections to cover:
- 8.1 Introduction
- 8.2 Simulation Waveforms (describe steady-state and transient waveforms)
- 8.3 Hardware Experimental Results
- 8.4 Performance Metrics Table (Metric | Target | Simulated | Experimental)
- 8.5 Efficiency and Power Loss Analysis
- 8.6 THD / Power Factor Analysis (if applicable)
- 8.7 Comparison with Existing Converters / Systems
- 8.8 Limitations and Future Improvements
Minimum: 1000 words.
""",
    },
    "mechanical": {
        "requirements": """
Write Chapter 3: Design Requirements & Specifications.
Sections to cover:
- 3.1 Introduction
- 3.2 Functional Requirements (minimum 10 numbered points)
- 3.3 Material Specifications and Selection Criteria
- 3.4 Dimensional and Geometric Constraints
- 3.5 Load, Stress, and Safety Requirements
- 3.6 Manufacturing Constraints and Tolerances
- 3.7 Quality and Safety Standards (IS, ASTM, ISO)
- 3.8 Economic Constraints and Budget
Minimum: 1200 words.
""",
        "system_design": """
Write Chapter 4: Design Methodology & CAD Modelling.
Sections to cover:
- 4.1 Design Approach and Methodology
- 4.2 Conceptual Design Alternatives Considered
- 4.3 Final Design Selection and Justification
- 4.4 Material Selection with Properties Table (Material | Yield Strength | Density | Cost/kg)
- 4.5 CAD Model Description (describe 3D components and assembly)
- 4.6 Engineering Drawing Description (views, critical dimensions, tolerances)
- 4.7 Bill of Materials
- 4.8 Design Standards Followed
Minimum: 1800 words.
""",
        "implementation": """
Write Chapter 5: Manufacturing & Fabrication.
Sections to cover:
- 5.1 Manufacturing Process Selection and Justification
- 5.2 Step-by-step Fabrication Sequence
- 5.3 Machining Operations and Parameters (cutting speed, feed, depth of cut)
- 5.4 Joining Methods (welding, fastening, adhesive bonding)
- 5.5 Surface Treatment and Finishing Processes
- 5.6 Quality Control Checks During Fabrication
- 5.7 Tools, Equipment, and Machines Used
- 5.8 Fabrication Challenges and Solutions
Minimum: 2500 words.
""",
        "database": """
Write Chapter 6: Structural & Thermal Analysis.
Sections to cover:
- 6.1 Analysis Methodology (Analytical / FEA / ANSYS)
- 6.2 Loading Conditions and Boundary Conditions
- 6.3 Static Structural Analysis (von Mises stress, strain, total deformation)
- 6.4 Factor of Safety Calculation
- 6.5 Thermal Analysis (heat transfer mode, temperature distribution)
- 6.6 Analysis Results Table (Component | Applied Load | Max Stress | Max Deformation | FOS)
- 6.7 Fatigue Life Estimation (if applicable)
- 6.8 Design Optimization Based on Analysis Results
Minimum: 1000 words. Include results table using pipe | separators.
""",
        "testing": """
Write Chapter 7: Testing & Validation.
Sections to cover:
- 7.1 Testing Plan and Strategy
- 7.2 Testing Equipment and Setup
- 7.3 Static Load Tests (Table: Load Case | Applied Load | Expected Deflection | Measured | Status — 15 rows)
- 7.4 Functional Performance Tests (10 test cases in same format)
- 7.5 Dimensional Accuracy Verification
- 7.6 Theoretical vs Experimental Comparison
- 7.7 Failure Mode and Root Cause Analysis
- 7.8 Corrective Actions Taken
Minimum: 1200 words.
""",
        "results": """
Write Chapter 8: Results & Performance Evaluation.
Sections to cover:
- 8.1 Introduction
- 8.2 Final Fabricated Prototype Description
- 8.3 Experimental vs Theoretical Comparison Table
- 8.4 Key Performance Metrics (load capacity, deflection, weight, efficiency)
- 8.5 FEA vs Physical Test Comparison
- 8.6 Cost Analysis and Budget Summary
- 8.7 Comparison with Existing Solutions
- 8.8 Limitations and Future Scope
Minimum: 1000 words.
""",
    },
    "civil": {
        "requirements": """
Write Chapter 3: Site Analysis & Design Requirements.
Sections to cover:
- 3.1 Introduction
- 3.2 Site Investigation and Soil Analysis
- 3.3 Load Requirements (dead load, live load, wind load, seismic load per IS 875 / IS 1893)
- 3.4 Design Standards and Codes Referenced (IS 456, IS 800, IS 1893, etc.)
- 3.5 Environmental and Sustainability Requirements
- 3.6 Regulatory and Statutory Approvals Required
- 3.7 Material Specifications
- 3.8 Project Constraints (budget, timeline, site limitations)
Minimum: 1200 words.
""",
        "system_design": """
Write Chapter 4: Structural System Design.
Sections to cover:
- 4.1 Structural System Selection and Justification (RCC / Steel / Composite)
- 4.2 Structural Layout, Grid, and Configuration
- 4.3 Load Path Analysis
- 4.4 Foundation System Selection (isolated / combined / raft / pile)
- 4.5 Analysis Software Used (STAAD Pro / ETABS / SAP2000) and Modelling Approach
- 4.6 Structural Model Description (members, supports, loading)
- 4.7 Design Philosophy (Limit State Design / Working Stress Method)
- 4.8 Design Assumptions
Minimum: 1800 words.
""",
        "implementation": """
Write Chapter 5: Structural Analysis & Calculations.
Sections to cover:
- 5.1 Load Calculations (dead, live, wind, seismic) with IS code formulas
- 5.2 Bending Moment and Shear Force Analysis
- 5.3 Beam Design Calculations (Limit State Method)
- 5.4 Column Design Calculations
- 5.5 Slab Design (one-way / two-way)
- 5.6 Foundation Design and Bearing Capacity Check
- 5.7 Computer-aided Analysis Results (STAAD / ETABS model outputs)
- 5.8 Design Summary Table
Minimum: 2500 words.
""",
        "database": """
Write Chapter 6: Detailed Design & Specifications.
Sections to cover:
- 6.1 Beam Reinforcement Details (Table: Member | Span | Size | Top Bars | Bottom Bars | Stirrups)
- 6.2 Column Reinforcement Details
- 6.3 Slab Reinforcement Details
- 6.4 Foundation Reinforcement Details
- 6.5 Connection Design (beam-column joints, base plate)
- 6.6 Bill of Materials (Table: Item | Specification | Quantity | Unit | Rate)
- 6.7 Construction Sequence
- 6.8 Quality Control and Inspection Specifications
Minimum: 1000 words. Include all tables using pipe | separators.
""",
        "testing": """
Write Chapter 7: Testing & Structural Validation.
Sections to cover:
- 7.1 Testing Methodology
- 7.2 Material Testing (concrete cube, steel tensile test — Table: Test ID | Sample | Expected | Result | Pass/Fail — 15 rows)
- 7.3 Load Testing (10 test cases: Member | Load | Expected Deflection | Actual | Status)
- 7.4 Non-Destructive Testing Methods Used
- 7.5 Software Model Validation
- 7.6 Comparison: Design Loads vs Actual Measured
- 7.7 Safety Factor Verification for All Members
- 7.8 Issues Found and Resolutions
Minimum: 1200 words.
""",
        "results": """
Write Chapter 8: Design Results & Safety Assessment.
Sections to cover:
- 8.1 Introduction
- 8.2 Structural Analysis Summary (max moments, shear, deflections)
- 8.3 Design Results Table (Member | Section Size | Reinforcement | Utilization Ratio)
- 8.4 Safety Factor Assessment for All Critical Members
- 8.5 Cost Estimation Summary
- 8.6 Comparison with Alternative Design
- 8.7 Environmental and Sustainability Impact
- 8.8 Limitations and Recommendations for Future Work
Minimum: 1000 words.
""",
    },
}


def _get_branch_key(branch: str) -> str:
    b = branch.lower()
    if "ece" in b or "extc" in b:
        return "ece"
    if "eee" in b or "electrical" in b:
        return "eee"
    if "mechanical" in b:
        return "mechanical"
    if "civil" in b:
        return "civil"
    return "default"


def get_branch_config(branch: str) -> dict:
    key     = _get_branch_key(branch)
    titles  = {
        "abstract"       : "Chapter 1: Abstract and Introduction",
        "literature"     : "Chapter 2: Literature Review",
        "requirements"   : "Chapter 3: Requirements Analysis",
        "system_design"  : "Chapter 4: System Design",
        "implementation" : "Chapter 5: Implementation",
        "database"       : "Chapter 6: Database Design",
        "testing"        : "Chapter 7: Testing",
        "results"        : "Chapter 8: Results and Discussion",
        "conclusion"     : "Chapter 9: Conclusion and Future Work",
        "references"     : "Chapter 10: References and Appendix",
    }
    prompts = dict(CHAPTER_PROMPTS)
    if key in _BRANCH_CHAPTER_TITLES:
        titles.update(_BRANCH_CHAPTER_TITLES[key])
    if key in _BRANCH_CHAPTER_PROMPTS:
        prompts.update(_BRANCH_CHAPTER_PROMPTS[key])
    return {"titles": titles, "prompts": prompts}


# ── State ─────────────────────────────────────────────────────────────────────

class ReportState(TypedDict):
    job             : dict
    chapter_plan    : dict
    context_summary : str
    chapters        : Dict[str, str]
    chapter_status  : Dict[str, str]
    retry_count     : Dict[str, int]
    quality_flags   : Dict[str, str]
    failed_chapters : List[str]
    docx_bytes      : bytes
    viva_content    : str
    summary_content : str
    drive_url       : str
    error           : str


# ── Node 1: Planner ───────────────────────────────────────────────────────────

async def planner_node(state: ReportState) -> ReportState:
    job = state["job"]
    branch  = job.get("domain", "CSE / IT")
    context = f"""
Project Title       : {job['title']}
Description         : {job['description']}
Engineering Branch  : {branch}
Tech Stack          : {job['tech_stack']}
College / University: {job['client']}
Student             : {job['student_name']}
Batch               : {job['batch_year']}
Key Modules/Features: {job.get('modules', 'As appropriate')}
"""
    r = await asyncio.to_thread(
        _get_client().messages.create,
        model      = "claude-haiku-4-5-20251001",
        max_tokens = 1024,
        messages   = [{
            "role"   : "user",
            "content": f"""
You are planning an academic project report.
Given this project, list what SPECIFIC things each chapter should mention.
Be project-specific — mention actual module names, table names, tech choices.

{context}

Return a JSON object with chapter names as keys and
1-2 sentences of specific guidance as values.
Chapters: abstract, literature, requirements, system_design,
implementation, database, testing, results, conclusion, references
"""
        }]
    )
    try:
        plan = json.loads(r.content[0].text)
    except Exception:
        plan = {ch: "" for ch in CHAPTER_NAMES}

    return {
        **state,
        "chapter_plan"    : plan,
        "context_summary" : context,
        "chapters"        : {},
        "chapter_status"  : {ch: "pending" for ch in CHAPTER_NAMES},
        "retry_count"     : {ch: 0 for ch in CHAPTER_NAMES},
        "quality_flags"   : {},
        "failed_chapters" : [],
    }


# ── Node 2: Chapter Generator ─────────────────────────────────────────────────

async def chapter_generator_node(state: ReportState) -> ReportState:
    context        = state["context_summary"]
    plan           = state["chapter_plan"]
    chapters       = dict(state.get("chapters", {}))
    status         = dict(state.get("chapter_status", {}))
    branch_prompts = get_branch_config(state["job"].get("domain", "CSE / IT"))["prompts"]

    def get_previous_summary(chapter_name: str) -> str:
        idx       = CHAPTER_NAMES.index(chapter_name)
        prev      = CHAPTER_NAMES[:idx]
        summaries = []
        for p in prev:
            if p in chapters:
                summaries.append(f"{p}: {chapters[p][:200]}...")
        return "\n".join(summaries) if summaries else ""

    async def gen_one(name: str) -> tuple:
        if status.get(name) == "done":
            return name, chapters.get(name, "")
        prev_summary = get_previous_summary(name)
        planner_note = plan.get(name, "")
        prompt = f"""
{context}

PLANNER GUIDANCE FOR THIS CHAPTER:
{planner_note}

PREVIOUSLY WRITTEN CHAPTERS (for consistency):
{prev_summary if prev_summary else "This is the first chapter."}

{branch_prompts[name]}

IMPORTANT:
- Be consistent with tech stack and module names from context
- Do not contradict anything in previous chapters
- Write in formal academic English
- No markdown formatting — plain text only
- Meet the minimum word count requirement
"""
        try:
            r = await asyncio.to_thread(
                _get_client().messages.create,
                model      = "claude-haiku-4-5-20251001",
                max_tokens = 2500,
                messages   = [{"role": "user", "content": prompt}]
            )
            return name, r.content[0].text if r.content else ""
        except Exception as e:
            print(f"[AGENT] Chapter '{name}' failed: {e}")
            return name, ""

    pending = [n for n in CHAPTER_NAMES if status.get(n) != "done"]
    for name in pending:
        _, content = await gen_one(name)
        chapters[name] = content   # update immediately so next chapter has context
        status[name]   = "generated"
    return {**state, "chapters": chapters, "chapter_status": status}


# ── Node 3: Quality Checker ───────────────────────────────────────────────────

def quality_node(state: ReportState) -> ReportState:
    chapters    = state["chapters"]
    status      = dict(state["chapter_status"])
    quality     = dict(state.get("quality_flags", {}))
    retry_count = dict(state.get("retry_count", {}))
    failed      = list(state.get("failed_chapters", []))

    MIN_WORDS = {
        "abstract": 800, "literature": 900, "requirements": 1000,
        "system_design": 1500, "implementation": 2000, "database": 800,
        "testing": 1000, "results": 800, "conclusion": 600, "references": 500
    }
    PLACEHOLDER_PHRASES = [
        "[insert", "lorem ipsum", "add content here",
        "to be filled", "tbd", "coming soon"
    ]
    for name in CHAPTER_NAMES:
        content    = chapters.get(name, "")
        word_count = len(content.split())
        min_words  = MIN_WORDS.get(name, 700)
        has_placeholder = any(p in content.lower() for p in PLACEHOLDER_PHRASES)
        if word_count < min_words or has_placeholder:
            retries = retry_count.get(name, 0)
            if retries < 1:
                status[name]      = "retry"
                quality[name]     = f"too_short:{word_count}/{min_words}"
                retry_count[name] = retries + 1
            else:
                status[name]  = "done"
                quality[name] = f"accepted_low:{word_count}"
                failed.append(name)
        else:
            status[name]  = "done"
            quality[name] = f"ok:{word_count}"

    return {
        **state,
        "chapter_status"  : status,
        "quality_flags"   : quality,
        "retry_count"     : retry_count,
        "failed_chapters" : failed,
    }


# ── Router ────────────────────────────────────────────────────────────────────

def route_after_quality(state: ReportState) -> Literal["retry", "assemble"]:
    if "retry" in state["chapter_status"].values():
        return "retry"
    return "assemble"


# ── Node 4: Retry ─────────────────────────────────────────────────────────────

async def retry_node(state: ReportState) -> ReportState:
    chapters       = dict(state["chapters"])
    status         = dict(state["chapter_status"])
    quality        = state["quality_flags"]
    context        = state["context_summary"]
    retry_names    = [n for n, s in status.items() if s == "retry"]
    branch_prompts = get_branch_config(state["job"].get("domain", "CSE / IT"))["prompts"]

    async def retry_one(name: str) -> tuple:
        reason = quality.get(name, "unknown")
        prompt = f"""
{context}

RETRY INSTRUCTION: Your previous attempt for this chapter was rejected.
Reason: {reason}

You MUST write significantly more content this time.
Be very detailed and comprehensive.
Include examples, detailed descriptions, and thorough explanations.

{branch_prompts[name]}

Previous attempt (expand significantly on this):
{chapters.get(name, '')[:500]}...
"""
        try:
            r = await asyncio.to_thread(
                _get_client().messages.create,
                model      = "claude-haiku-4-5-20251001",
                max_tokens = 2500,
                messages   = [{"role": "user", "content": prompt}]
            )
            return name, r.content[0].text if r.content else ""
        except Exception as e:
            print(f"[AGENT] Retry '{name}' failed: {e}")
            return name, chapters.get(name, "")  # keep previous attempt on failure

    for name in retry_names:
        _, content = await retry_one(name)
        chapters[name] = content
        status[name]   = "generated"
    return {**state, "chapters": chapters, "chapter_status": status}


# ── Diagram generators ───────────────────────────────────────────────────────

def _make_diagram(draw_fn, figsize=(10, 6)) -> io.BytesIO:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis("off")
    draw_fn(ax)
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _parse_modules(job) -> list:
    """Return a clean list of module/feature names from the job."""
    text = job.get("modules", "")
    if not text:
        return ["User Login", "Dashboard", "Main Feature", "Reports", "Admin Panel"]
    items = []
    for sep in ["\n", ",", ";", "|"]:
        if sep in text:
            items = [m.strip() for m in text.split(sep) if m.strip()]
            break
    if not items:
        items = [text.strip()]
    cleaned = []
    for item in items:
        item = item.lstrip("0123456789.-) ").strip()
        if item and len(item) > 1:
            cleaned.append(item[:38])
    return cleaned[:10] if cleaned else ["Login", "Dashboard", "Data Management", "Reports"]


def _parse_tech(job) -> list:
    """Return tech stack as a list."""
    tech = job.get("tech_stack", "")
    return [t.strip() for t in tech.split(",") if t.strip()][:8]


def _extract_entities(job) -> list:
    """
    Derive realistic entity names from modules + title.
    Returns list of (EntityName, [field1, field2, ...]) tuples.
    """
    modules  = _parse_modules(job)
    title    = job.get("title", "Project")

    # Try to extract nouns from module names as entity candidates
    stop = {"management", "system", "module", "panel", "page", "feature",
            "handling", "processing", "and", "the", "for", "of"}
    candidates = []
    for m in modules:
        for word in m.replace("-", " ").split():
            w = word.strip("()").lower()
            if len(w) > 2 and w not in stop:
                candidates.append(w.capitalize())

    # Deduplicate preserving order
    seen = set()
    entities_raw = []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            entities_raw.append(c)

    # Always include User
    if "User" not in seen:
        entities_raw.insert(0, "User")

    # Pick 4 entities
    chosen = entities_raw[:4]
    if len(chosen) < 3:
        chosen = ["User", "Admin", title.split()[0]][:4]

    # Build field lists per entity
    generic_pk = lambda e: f"{e.lower()}_id (PK)"
    result = []
    for e in chosen:
        el = e.lower()
        if el in ("user", "student", "member", "customer", "employee"):
            fields = [generic_pk(e), "name", "email", "password", "created_at"]
        elif el in ("admin", "administrator", "manager"):
            fields = [generic_pk(e), "username", "email", "role", "last_login"]
        elif el in ("product", "item", "book", "course", "article"):
            fields = [generic_pk(e), "title", "description", "price", "stock"]
        elif el in ("order", "booking", "reservation", "transaction", "purchase"):
            fields = [generic_pk(e), "user_id (FK)", "amount", "status", "date"]
        elif el in ("category", "department", "branch", "type"):
            fields = [generic_pk(e), "name", "description"]
        elif el in ("payment", "fee", "invoice"):
            fields = [generic_pk(e), "user_id (FK)", "amount", "method", "date"]
        elif el in ("report", "log", "history", "record"):
            fields = [generic_pk(e), "user_id (FK)", "content", "created_at"]
        elif el in ("notification", "alert", "message"):
            fields = [generic_pk(e), "user_id (FK)", "text", "is_read", "sent_at"]
        else:
            fields = [generic_pk(e), "name", "description", "status", "created_at"]
        result.append((e, fields))
    return result


def _draw_actor(ax, x, y_center, label, color="#6C63FF"):
    """Draw a UML stick figure actor."""
    import matplotlib.pyplot as plt
    ax.add_patch(plt.Circle((x, y_center + 0.5), 0.35, color=color, zorder=3))
    ax.plot([x, x],            [y_center+0.15, y_center-0.7],  color="#333", lw=1.8)
    ax.plot([x-0.55, x+0.55],  [y_center-0.15, y_center-0.15], color="#333", lw=1.8)
    ax.plot([x, x-0.45],       [y_center-0.7,  y_center-1.3],  color="#333", lw=1.8)
    ax.plot([x, x+0.45],       [y_center-0.7,  y_center-1.3],  color="#333", lw=1.8)
    ax.text(x, y_center-1.6, label, ha="center", fontsize=8.5, fontweight="bold")


def _diagram_architecture(job) -> io.BytesIO:
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.pyplot as plt

    tech   = _parse_tech(job)
    title  = job.get("title", "System")

    # Assign tech to layers heuristically
    frontend_kw  = {"react","vue","angular","flutter","html","css","javascript","typescript","nextjs","svelte"}
    backend_kw   = {"flask","django","fastapi","express","spring","node","nodejs","laravel","rails","asp"}
    db_kw        = {"mysql","postgresql","mongodb","sqlite","firebase","redis","oracle","cassandra","dynamodb"}
    infra_kw     = {"docker","kubernetes","aws","gcp","azure","nginx","apache","heroku","vercel","firebase"}

    layers = {"Frontend": [], "Backend / API": [], "Database": [], "Infrastructure": []}
    for t in tech:
        tl = t.lower()
        if tl in frontend_kw:   layers["Frontend"].append(t)
        elif tl in backend_kw:  layers["Backend / API"].append(t)
        elif tl in db_kw:       layers["Database"].append(t)
        elif tl in infra_kw:    layers["Infrastructure"].append(t)
        else:
            # Assign to least-filled layer
            least = min(layers, key=lambda k: len(layers[k]))
            layers[least].append(t)

    # Fill empty layers with generic labels
    defaults = {"Frontend": ["Web / Mobile UI"], "Backend / API": ["Business Logic"],
                "Database": ["Data Store"], "Infrastructure": ["Cloud Platform"]}
    for k in layers:
        if not layers[k]:
            layers[k] = defaults[k]

    layer_colors = {
        "Frontend":        "#DBEAFE",
        "Backend / API":   "#DCFCE7",
        "Database":        "#FEF9C3",
        "Infrastructure":  "#FCE7F3",
    }

    def draw(ax):
        layer_names = list(layers.keys())
        ys = [8.5, 6.5, 4.5, 2.5]
        for (name, y) in zip(layer_names, ys):
            color = layer_colors[name]
            ax.add_patch(FancyBboxPatch((0.5, y - 0.75), 9, 1.5,
                         boxstyle="round,pad=0.15",
                         facecolor=color, edgecolor="#555", lw=1.5))
            ax.text(2.0, y, name, ha="center", va="center",
                    fontsize=10, fontweight="bold", color="#222")
            components = ", ".join(layers[name])
            ax.text(6.0, y, components, ha="center", va="center",
                    fontsize=9, color="#444",
                    bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="#aaa", lw=0.8))

        for y1, y2 in zip([7.75, 5.75, 3.75], [7.25, 5.25, 3.25]):
            ax.annotate("", xy=(5, y2), xytext=(5, y1),
                        arrowprops=dict(arrowstyle="<->", color="#555", lw=1.5))

        ax.text(5, 9.6, f"System Architecture — {title[:45]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(11, 8))


def _diagram_usecase(job) -> io.BytesIO:
    from matplotlib.patches import Ellipse, FancyBboxPatch
    import matplotlib.pyplot as plt

    modules = _parse_modules(job)
    title   = job.get("title", "System")

    # Classify modules: admin vs user
    admin_kw = {"admin", "manage", "delete", "approve", "configure", "setting",
                "dashboard", "monitor", "report", "user management"}
    user_ucs  = []
    admin_ucs = []
    for m in modules:
        if any(kw in m.lower() for kw in admin_kw):
            admin_ucs.append(m)
        else:
            user_ucs.append(m)
    if len(user_ucs) < 3:
        user_ucs  = modules[:6]
        admin_ucs = modules[6:]

    user_ucs  = user_ucs[:6]
    admin_ucs = admin_ucs[:4]
    all_ucs   = user_ucs + admin_ucs

    def draw(ax):
        n  = len(all_ucs)
        y0 = 9.2
        gap = 8.4 / max(n - 1, 1)
        ys = [y0 - i * gap for i in range(n)]

        y_min, y_max = min(ys) - 0.6, max(ys) + 0.6

        # System boundary
        ax.add_patch(FancyBboxPatch((2.8, y_min), 6.8, y_max - y_min,
                     boxstyle="round,pad=0.1", fill=False,
                     edgecolor="#888", lw=1.5, linestyle="--"))
        ax.text(6.2, y_max + 0.3, f"«system»  {title[:30]}",
                ha="center", fontsize=8.5, color="#555", style="italic")

        # Use cases
        for i, (uc, y) in enumerate(zip(all_ucs, ys)):
            is_admin = i >= len(user_ucs)
            fc = "#FEF9C3" if is_admin else "#DBEAFE"
            ec = "#CA8A04" if is_admin else "#2563EB"
            el = Ellipse((6.2, y), width=6.0, height=0.72,
                         facecolor=fc, edgecolor=ec, lw=1.2)
            ax.add_patch(el)
            ax.text(6.2, y, uc[:34], ha="center", va="center", fontsize=8.5)

        # User actor (left)
        _draw_actor(ax, 1.1, 5.5, "User", color="#6C63FF")
        for y in ys[:len(user_ucs)]:
            ax.annotate("", xy=(2.8, y), xytext=(1.4, 5.5),
                        arrowprops=dict(arrowstyle="-", color="#555", lw=0.9))

        # Admin actor (right) if admin use cases exist
        if admin_ucs:
            _draw_actor(ax, 9.1, 5.5, "Admin", color="#F59E0B")
            for y in ys[len(user_ucs):]:
                ax.annotate("", xy=(9.2, y), xytext=(8.8 + (10 - 9.2) * 0.1, 5.5),
                            arrowprops=dict(arrowstyle="-", color="#888", lw=0.9))

        ax.text(5, 10.3, f"Use Case Diagram — {title[:40]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(11, 10))


def _diagram_flowchart(job) -> io.BytesIO:
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.pyplot as plt

    modules = _parse_modules(job)
    title   = job.get("title", "System")

    # Build flow steps: Start → Login → [core modules] → End
    steps = []
    steps.append(("Start", "#6C63FF", "white"))
    steps.append(("User Login / Authentication", "#DBEAFE", "#111"))
    for m in modules[:5]:
        steps.append((m, "#DCFCE7", "#111"))
    steps.append(("Save / Update Records", "#FEF9C3", "#111"))
    steps.append(("Generate Report / Output", "#DBEAFE", "#111"))
    steps.append(("End", "#6C63FF", "white"))

    def draw(ax):
        n    = len(steps)
        y0   = 9.5
        gap  = 9.0 / max(n - 1, 1)
        ys   = [y0 - i * gap for i in range(n)]
        x    = 5.0
        w, h = 5.5, 0.65

        for (label, fc, tc), y in zip(steps, ys):
            ax.add_patch(FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle="round,pad=0.12",
                         facecolor=fc, edgecolor="#555", lw=1.2))
            ax.text(x, y, label[:40], ha="center", va="center",
                    fontsize=8.5, fontweight="bold", color=tc)

        for i in range(len(steps) - 1):
            y1 = ys[i]   - h/2
            y2 = ys[i+1] + h/2
            ax.annotate("", xy=(x, y2), xytext=(x, y1),
                        arrowprops=dict(arrowstyle="->", color="#555", lw=1.4))

        # Validation error branch from step 1 (login)
        y_val = ys[1]
        ax.annotate("", xy=(8.8, y_val), xytext=(x + w/2, y_val),
                    arrowprops=dict(arrowstyle="->", color="#DC2626", lw=1.2))
        ax.text(9.0, y_val + 0.25, "Error /\nInvalid",
                ha="center", fontsize=7.5, color="#DC2626")

        ax.text(x, 10.2, f"System Flow Diagram — {title[:40]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(9, 12))


def _diagram_dfd(job) -> io.BytesIO:
    from matplotlib.patches import FancyBboxPatch, Circle
    import matplotlib.pyplot as plt

    modules = _parse_modules(job)
    title   = job.get("title", "System")

    # Pick up to 3 internal processes from modules
    processes = [m[:22] for m in modules[:3]]
    if not processes:
        processes = ["Process Data", "Store Records", "Generate Output"]

    def draw(ax):
        # External entities
        for (label, x, y) in [("User", 0.5, 7.5), ("Admin", 0.5, 2.5)]:
            ax.add_patch(FancyBboxPatch((x, y - 0.4), 1.6, 0.8,
                         boxstyle="square,pad=0.1",
                         facecolor="#FEF9C3", edgecolor="#555", lw=1.5))
            ax.text(x + 0.8, y, label, ha="center", va="center",
                    fontsize=9, fontweight="bold")

        # Central system circle
        ax.add_patch(Circle((5, 5), 1.6, facecolor="#DCFCE7",
                            edgecolor="#16A34A", lw=2))
        name_parts = (title[:12] + "\nSystem").split("\n")
        ax.text(5, 5.2, name_parts[0], ha="center", fontsize=8.5, fontweight="bold")
        ax.text(5, 4.8, name_parts[1] if len(name_parts) > 1 else "", ha="center", fontsize=8.5)

        # Internal process bubbles
        proc_xs = [4.0, 5.0, 6.0]
        proc_ys = [7.5, 8.5, 7.5] if len(processes) > 2 else [4.5, 5.5]
        for i, (proc, px, py) in enumerate(zip(processes,
                                                [3.8, 5, 6.2],
                                                [8.2, 8.8, 8.2])):
            ax.add_patch(Circle((px, py), 0.75,
                         facecolor="#DBEAFE", edgecolor="#2563EB", lw=1.2))
            ax.text(px, py, proc[:14], ha="center", va="center", fontsize=7.5,
                    wrap=True)
            ax.annotate("", xy=(px, py - 0.75), xytext=(5, 6.6),
                        arrowprops=dict(arrowstyle="<->", color="#333", lw=1.2))

        # Data stores (right side)
        stores = [("Database", 8.8, 7.0), ("File Storage", 8.8, 3.5)]
        for label, sx, sy in stores:
            ax.add_patch(FancyBboxPatch((sx - 1.0, sy - 0.35), 2.0, 0.7,
                         boxstyle="round,pad=0.08",
                         facecolor="#EDE9FE", edgecolor="#7C3AED", lw=1.2))
            ax.text(sx, sy, label, ha="center", va="center", fontsize=8.5)
            ax.annotate("", xy=(sx - 1.0, sy), xytext=(6.6, 5.2),
                        arrowprops=dict(arrowstyle="<->", color="#7C3AED", lw=1.1))

        # Arrows from external entities to system
        for (ex, ey), label in [((2.1, 7.4), "Input Data"), ((2.1, 2.6), "Admin Request")]:
            ax.annotate("", xy=(3.5, 5.2), xytext=(ex, ey),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
            ax.text((2.1 + 3.5)/2 - 0.2, (ey + 5.2)/2, label,
                    ha="center", fontsize=7.5, color="#555")
        # Response arrow back
        ax.annotate("", xy=(2.1, 7.0), xytext=(3.5, 4.8),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1.2))
        ax.text(2.7, 5.8, "Output /\nStatus", ha="center", fontsize=7.5, color="#555")

        ax.text(5, 9.7, f"Data Flow Diagram (L0) — {title[:35]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(11, 9))


def _diagram_er(job) -> io.BytesIO:
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.pyplot as plt

    entities = _extract_entities(job)
    title    = job.get("title", "System")
    n        = len(entities)

    def draw(ax):
        # Position entities: up to 4 in a row
        positions = [(2.0, 7.0), (6.5, 7.0), (2.0, 3.0), (6.5, 3.0)][:n]
        box_w, row_h = 2.6, 0.42

        placed = {}
        for (ename, fields), (cx, cy) in zip(entities, positions):
            h = (len(fields) + 1) * row_h
            # Header bar
            ax.add_patch(FancyBboxPatch((cx - box_w/2, cy - row_h), box_w, row_h,
                         facecolor="#2563EB", edgecolor="#2563EB",
                         boxstyle="round,pad=0.03", lw=0))
            ax.text(cx, cy - row_h/2, ename, ha="center", va="center",
                    fontsize=10, fontweight="bold", color="white")
            # Fields
            ax.add_patch(FancyBboxPatch((cx - box_w/2, cy - h), box_w, h - row_h,
                         facecolor="#DBEAFE", edgecolor="#2563EB",
                         boxstyle="round,pad=0.03", lw=1.2))
            for i, field in enumerate(fields):
                is_pk = "(PK)" in field
                is_fk = "(FK)" in field
                fc = "#1D4ED8" if is_pk else ("#7C3AED" if is_fk else "#111")
                ax.text(cx, cy - row_h * (i + 1.5), ("🔑 " if is_pk else "  ") + field,
                        ha="center", va="center", fontsize=7.5, color=fc)
            placed[ename] = (cx, cy - row_h)

        # Draw relationship lines between consecutive entities
        entity_names = [e[0] for e in entities]
        rel_labels   = ["1..N", "1..1", "1..N", "1..1"]
        pairs = [(entity_names[i], entity_names[i+1], rel_labels[i])
                 for i in range(len(entity_names) - 1)]
        for e1, e2, lbl in pairs:
            if e1 in placed and e2 in placed:
                x1, y1 = placed[e1]
                x2, y2 = placed[e2]
                mx, my = (x1+x2)/2, (y1+y2)/2
                ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                            arrowprops=dict(arrowstyle="-", color="#333", lw=1.5))
                ax.text(mx + 0.15, my + 0.2, lbl, ha="center", fontsize=8,
                        color="#555",
                        bbox=dict(fc="white", ec="none", pad=1))

        ax.text(4.5, 9.5, f"ER Diagram — {title[:40]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(11, 9))


def _diagram_module_interaction(job) -> io.BytesIO:
    """Component/module interaction diagram showing how project modules connect."""
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.pyplot as plt

    modules = _parse_modules(job)
    title   = job.get("title", "System")

    # Central hub + surrounding modules
    hub     = title[:18]
    spokes  = modules[:6]
    n       = len(spokes)

    def draw(ax):
        import math
        cx, cy, r = 5.0, 5.0, 3.2

        # Central box
        ax.add_patch(FancyBboxPatch((cx - 1.4, cy - 0.5), 2.8, 1.0,
                     boxstyle="round,pad=0.15",
                     facecolor="#6C63FF", edgecolor="#4338CA", lw=2))
        ax.text(cx, cy, hub, ha="center", va="center",
                fontsize=10, fontweight="bold", color="white")

        colors = ["#DBEAFE", "#DCFCE7", "#FEF9C3", "#FCE7F3", "#E0F2FE", "#F3F4F6"]
        for i, mod in enumerate(spokes):
            angle  = 2 * math.pi * i / n - math.pi / 2
            mx     = cx + r * math.cos(angle)
            my     = cy + r * math.sin(angle)
            ax.add_patch(FancyBboxPatch((mx - 1.3, my - 0.38), 2.6, 0.76,
                         boxstyle="round,pad=0.1",
                         facecolor=colors[i % len(colors)],
                         edgecolor="#555", lw=1.2))
            ax.text(mx, my, mod[:24], ha="center", va="center", fontsize=8)
            # Arrow from hub to module
            dx = mx - cx; dy = my - cy
            dist = math.hypot(dx, dy)
            ax.annotate("", xy=(cx + dx/dist*1.5, cy + dy/dist*0.55),
                        xytext=(cx + dx/dist*(r - 1.4), cy + dy/dist*(r - 0.42)),
                        arrowprops=dict(arrowstyle="<->", color="#555", lw=1.2))

        ax.text(cx, 9.6, f"Module Interaction — {title[:40]}",
                ha="center", fontsize=11, fontweight="bold")

    return _make_diagram(draw, figsize=(11, 9))


# ── Node 5: Assembler ─────────────────────────────────────────────────────────

def assembler_node(state: ReportState) -> ReportState:
    """
    Assembles all 10 chapters into a properly formatted .docx.
    Applies heading colors, Times New Roman body, standard margins,
    and converts pipe-delimited text blocks to real Word tables.
    """
    job      = state["job"]
    chapters = state["chapters"]

    from docx.shared      import Pt, Cm, RGBColor
    from docx.enum.text   import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns     import qn
    from docx.oxml        import OxmlElement

    doc_color_hex = job.get("doc_color_hex", "#1A4B8C")
    heading_rgb   = _hex_to_rgb(doc_color_hex)
    h2_rgb        = _lighten_rgb(heading_rgb, 0.25)
    accent_hex    = doc_color_hex.lstrip("#")

    doc = Document()

    # ── Page margins & header/footer ──────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
        _add_header_footer(section, job, heading_rgb, accent_hex)

    # ── Helpers ────────────────────────────────────────────────────
    def _shade_para(para, fill_hex):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
        pPr.append(shd)

    def _set_line_spacing(para, spacing=1.15):
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = Pt(12 * spacing)

    def _left_border(para, color_hex):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "thick")
        left.set(qn("w:sz"),    "24")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), color_hex.lstrip("#"))
        pBdr.append(left)
        pPr.append(pBdr)

    def _set_font(run, name="Calibri", size=11, bold=False, color=None):
        run.font.name  = name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color

    def add_centered(text, bold=False, size=11, font="Calibri", color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, font, size, bold, color)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_font(run, "Calibri", 11)
        _set_line_spacing(p, 1.3)
        p.paragraph_format.space_after = Pt(6)
        return p

    # ── Global styles ──────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1_style = doc.styles["Heading 1"]
    h1_style.font.name  = "Calibri"
    h1_style.font.bold  = True
    h1_style.font.size  = Pt(20)
    h1_style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    h1_style.paragraph_format.space_before = Pt(0)
    h1_style.paragraph_format.space_after  = Pt(12)

    h2_style = doc.styles["Heading 2"]
    h2_style.font.name  = "Calibri"
    h2_style.font.bold  = True
    h2_style.font.size  = Pt(14)
    h2_style.font.color.rgb = heading_rgb
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after  = Pt(4)

    h3_style = doc.styles["Heading 3"]
    h3_style.font.name   = "Calibri"
    h3_style.font.bold   = False
    h3_style.font.italic = True
    h3_style.font.size   = Pt(12)
    h3_style.font.color.rgb = h2_rgb
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after  = Pt(2)

    # ── Cover page ─────────────────────────────────────────────────
    # Top accent bar
    bar = doc.add_paragraph()
    _shade_para(bar, doc_color_hex)
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    run = bar.add_run(f"  {job.get('client', 'Academic Project Report')}  ·  Academic Project Report")
    _set_font(run, "Calibri", 10, color=RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(job["title"])
    _set_font(t_run, "Calibri", 26, bold=True, color=heading_rgb)

    # Divider after title
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(6)
    div.paragraph_format.space_after  = Pt(6)
    _div_pPr = div._p.get_or_add_pPr()
    _div_bdr = OxmlElement('w:pBdr')
    _div_bot = OxmlElement('w:bottom')
    _div_bot.set(qn('w:val'),   'single')
    _div_bot.set(qn('w:sz'),    '12')
    _div_bot.set(qn('w:space'), '1')
    _div_bot.set(qn('w:color'), accent_hex)
    _div_bdr.append(_div_bot)
    _div_pPr.append(_div_bdr)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "A Project Report submitted in partial fulfillment of the\n"
        "requirements for the degree of Bachelor of Engineering"
    )
    _set_font(sub_run, "Calibri", 11, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph("")

    # Info card — 2-column borderless table
    info_rows = [
        ("Submitted by",    job.get("student_name", "")),
        ("Guide",           job.get("guider_name", "")),
        ("Semester / Class",job.get("semester", "")),
        ("Batch",           str(job.get("batch_year", ""))),
        ("Institution",     job.get("client", "")),
        ("Branch / Domain", job.get("domain", "CSE / IT")),
    ]
    info_tbl = doc.add_table(rows=len(info_rows), cols=2)
    info_tbl.style = "Table Grid"
    _itblPr  = info_tbl._tbl.tblPr
    _itblBdr = OxmlElement('w:tblBorders')
    for _side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        _b = OxmlElement(f'w:{_side}')
        _b.set(qn('w:val'), 'none')
        _itblBdr.append(_b)
    _itblPr.append(_itblBdr)
    _itblJc = OxmlElement('w:jc')
    _itblJc.set(qn('w:val'), 'center')
    _itblPr.append(_itblJc)
    for r_idx, (label, value) in enumerate(info_rows):
        lbl_cell = info_tbl.rows[r_idx].cells[0]
        val_cell = info_tbl.rows[r_idx].cells[1]
        lbl_cell.text = ""
        val_cell.text = ""
        _shade_para(lbl_cell.paragraphs[0], "F3F4F6")
        _shade_para(val_cell.paragraphs[0], "F3F4F6")
        lbl_run = lbl_cell.paragraphs[0].add_run(label)
        _set_font(lbl_run, "Calibri", 11, bold=True, color=heading_rgb)
        val_run = val_cell.paragraphs[0].add_run(value)
        _set_font(val_run, "Calibri", 11)
        for cell in (lbl_cell, val_cell):
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(4)

    doc.add_paragraph("")
    # Tech stack chips line
    tech = job.get("tech_stack", "")
    if tech:
        chips_p = doc.add_paragraph()
        chips_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for chip in [t.strip() for t in tech.split(",")][:6]:
            r = chips_p.add_run(f"  {chip}  ")
            _set_font(r, "Calibri", 9, color=heading_rgb)

    # Rule above bottom bar
    doc.add_paragraph("")
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(4)
    rule_p.paragraph_format.space_after  = Pt(2)
    _rule_pPr = rule_p._p.get_or_add_pPr()
    _rule_bdr = OxmlElement('w:pBdr')
    _rule_bot = OxmlElement('w:bottom')
    _rule_bot.set(qn('w:val'),   'single')
    _rule_bot.set(qn('w:sz'),    '6')
    _rule_bot.set(qn('w:space'), '1')
    _rule_bot.set(qn('w:color'), accent_hex)
    _rule_bdr.append(_rule_bot)
    _rule_pPr.append(_rule_bdr)

    # Bottom bar
    bot = doc.add_paragraph()
    _shade_para(bot, doc_color_hex)
    bot_run = bot.add_run(f"  {job.get('domain','CSE / IT')}  ·  {job['batch_year']}  ·  {job.get('client','')}")
    _set_font(bot_run, "Calibri", 9, color=RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────
    toc_h = doc.add_heading("Table of Contents", 1)
    _shade_para(toc_h, doc_color_hex)
    for run in toc_h.runs:
        _set_font(run, "Calibri", 20, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    toc_para = doc.add_paragraph()
    toc_run_begin = toc_para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    toc_run_begin._r.append(fld_begin)

    toc_run_instr = toc_para.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    toc_run_instr._r.append(instr_text)

    toc_run_sep = toc_para.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    toc_run_sep._r.append(fld_sep)

    toc_run_placeholder = toc_para.add_run()
    toc_run_placeholder.text = '[Update this field in Word: right-click → Update Field]'
    toc_run_placeholder.font.size = Pt(10)
    toc_run_placeholder.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    toc_run_end = toc_para.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    toc_run_end._r.append(fld_end)

    doc.add_page_break()

    chapter_titles = get_branch_config(
        job.get("domain", "CSE / IT")
    )["titles"]

    _CHAPTER_DIAGRAMS = {
        "requirements" : [("Use Case Diagram",           _diagram_usecase)],
        "system_design": [("System Architecture Diagram",_diagram_architecture),
                          ("System Flow Diagram",        _diagram_flowchart),
                          ("Data Flow Diagram (L0)",     _diagram_dfd),
                          ("Module Interaction Diagram", _diagram_module_interaction)],
        "database"     : [("Entity Relationship Diagram",_diagram_er)],
    }

    for key, title in chapter_titles.items():
        doc.add_page_break()
        h1 = doc.add_heading(title, 1)
        _shade_para(h1, doc_color_hex)
        for run in h1.runs:
            _set_font(run, "Calibri", 20, bold=True,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
        doc.add_page_break()

        content = chapters.get(key, "Content not available.")

        for para in content.split("\n\n"):
            if not para.strip():
                continue
            stripped = para.strip()
            lines    = [l.strip() for l in stripped.splitlines() if l.strip()]

            # ── Detect pipe-delimited table ────────────────────────
            pipe_lines = [l for l in lines if '|' in l]
            if len(pipe_lines) >= 2 and len(pipe_lines) >= len(lines) * 0.6:
                rows = []
                for line in lines:
                    if set(line.replace(' ', '')) <= set('|-+'):
                        continue
                    cells = [c.strip() for c in line.strip('|').split('|')]
                    if any(c for c in cells):
                        rows.append(cells)
                if rows:
                    max_cols = max(len(r) for r in rows)
                    rows = [r + [''] * (max_cols - len(r)) for r in rows]
                    tbl = doc.add_table(rows=len(rows), cols=max_cols)
                    tbl.style  = 'Table Grid'
                    tbl.autofit = False
                    _tblPr = tbl._tbl.tblPr
                    _tblW  = OxmlElement('w:tblW')
                    _tblW.set(qn('w:w'),    '5000')
                    _tblW.set(qn('w:type'), 'pct')
                    _tblPr.append(_tblW)
                    _tblCM = OxmlElement('w:tblCellMar')
                    for _side, _val in (('top','60'),('bottom','60'),('left','80'),('right','80')):
                        _m = OxmlElement(f'w:{_side}')
                        _m.set(qn('w:w'),    _val)
                        _m.set(qn('w:type'), 'dxa')
                        _tblCM.append(_m)
                    _tblPr.append(_tblCM)
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            cell = tbl.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(cell_text)
                            if r_idx == 0:
                                _set_font(run, "Calibri", 11, bold=True,
                                          color=RGBColor(0xFF, 0xFF, 0xFF))
                                _shade_para(cell.paragraphs[0], doc_color_hex)
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                _set_font(run, "Calibri", 10)
                                if r_idx % 2 == 0:
                                    _shade_para(cell.paragraphs[0], "F3F4F6")
                    doc.add_paragraph("")
                continue

            # ── Detect subheadings (short lines starting with digit) ──
            if (len(stripped) < 80
                    and not stripped.endswith(".")
                    and stripped[0].isdigit()):
                if stripped.count('.') >= 2:
                    h3 = doc.add_heading(stripped, 3)
                    for run in h3.runs:
                        _set_font(run, "Calibri", 12, bold=False, color=h2_rgb)
                        run.font.italic = True
                else:
                    h2 = doc.add_heading(stripped, 2)
                    _left_border(h2, doc_color_hex)
                    for run in h2.runs:
                        _set_font(run, "Calibri", 14, bold=True, color=h2_rgb)
            else:
                add_body(stripped)

        # ── Inject diagrams for specific chapters ─────────────────
        if key in _CHAPTER_DIAGRAMS:
            for diag_title, diag_fn in _CHAPTER_DIAGRAMS[key]:
                try:
                    doc.add_page_break()
                    dh = doc.add_heading(diag_title, 2)
                    for run in dh.runs:
                        run.font.color.rgb = h2_rgb
                    buf_img = diag_fn(job)
                    doc.add_picture(buf_img, width=Cm(15))
                    doc.add_paragraph("")
                except Exception as diag_err:
                    print(f"[DIAGRAM] {diag_title} failed: {diag_err}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return {**state, "docx_bytes": buf.getvalue()}


# ── Node 6: Appendix (Viva + Executive Summary) ───────────────────────────────

async def appendix_node(state: ReportState) -> ReportState:
    """
    Generates Appendix C (Executive Summary) and Appendix D (Viva Questions)
    via two parallel Haiku calls, then appends them to the existing docx bytes.
    """
    context       = state["context_summary"]
    job           = state["job"]
    doc_color_hex = job.get("doc_color_hex", "#1A4B8C")
    heading_rgb   = _hex_to_rgb(doc_color_hex)

    async def gen_summary() -> str:
        r = await asyncio.to_thread(
            _get_client().messages.create,
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 1024,
            messages   = [{
                "role"   : "user",
                "content": f"""
{context}

Write a 1-page standalone Executive Summary for this academic project report.
It must be self-contained — a reader with no prior context should understand
the project, its objectives, methodology, key results, and conclusion.
Cover: purpose, scope, approach, key findings, conclusion.
Approximately 400-500 words. No markdown. Plain academic English.
"""
            }]
        )
        return r.content[0].text

    async def gen_viva() -> str:
        r = await asyncio.to_thread(
            _get_client().messages.create,
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 2048,
            messages   = [{
                "role"   : "user",
                "content": f"""
{context}

Generate exactly 25 likely viva voce / oral examination question-and-answer
pairs for this academic project.
Cover: domain fundamentals, tech stack choices, design decisions,
specific module logic, testing approach, and future work.
Format strictly as:
Q1: <question>
A1: <answer>

Q2: <question>
A2: <answer>

...up to Q25/A25. No markdown. Plain text only.
"""
            }]
        )
        return r.content[0].text

    summary_text, viva_text = await asyncio.gather(gen_summary(), gen_viva())

    # ── Append to existing docx ────────────────────────────────────
    buf = io.BytesIO(state["docx_bytes"])
    doc = Document(buf)

    _add_chapter_separator(doc)
    h1 = doc.add_heading("Appendix C: Executive Summary", 1)
    for run in h1.runs:
        run.font.color.rgb = heading_rgb
    for para in summary_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    _add_chapter_separator(doc)
    h1 = doc.add_heading("Appendix D: Viva Questions", 1)
    for run in h1.runs:
        run.font.color.rgb = heading_rgb
    for para in viva_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    return {
        **state,
        "docx_bytes"     : out.getvalue(),
        "viva_content"   : viva_text,
        "summary_content": summary_text,
    }


# ── Node 7: Delivery ──────────────────────────────────────────────────────────

def delivery_node(state: ReportState) -> ReportState:
    """
    Packages the .docx into a ZIP, uploads to Firebase Storage,
    generates a 24-hour signed download URL, updates Firestore, sends FCM.
    No Google Drive dependency.
    """
    from firebase_admin import firestore, storage, messaging

    job  = state["job"]
    jid  = job["job_id"]
    db   = firestore.client()
    blob = None

    try:
        # ── Build ZIP in memory ────────────────────────────────────
        zip_buf  = io.BytesIO()
        filename = f"{job['title']}_Project_Report.docx".replace("/", "-")
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(filename, state["docx_bytes"])
        zip_buf.seek(0)
        zip_bytes = zip_buf.getvalue()

        # ── Upload ZIP to Firebase Storage ────────────────────────
        import uuid, urllib.parse
        bucket     = storage.bucket()
        zip_name   = f"{job['title']}_Report.zip".replace("/", "-")
        blob       = bucket.blob(f"reports/{jid}/{zip_name}")
        blob.upload_from_string(zip_bytes, content_type="application/zip", timeout=120)

        # ── Generate Firebase Storage download URL ────────────────
        download_token = str(uuid.uuid4())
        blob.metadata  = {"firebaseStorageDownloadTokens": download_token}
        blob.patch()
        encoded_path = urllib.parse.quote(f"reports/{jid}/{zip_name}", safe="")
        bucket_name  = blob.bucket.name
        download_url = (
            f"https://firebasestorage.googleapis.com/v0/b/{bucket_name}"
            f"/o/{encoded_path}?alt=media&token={download_token}"
        )
        expires_at = datetime.now(tz=timezone.utc) + timedelta(hours=24)

        # ── Update Firestore ──────────────────────────────────────
        db.collection("jobs").document(jid).update({
            "status"       : "done",
            "download_url" : download_url,
            "expires_at"   : expires_at,
            "pages"        : 94,
            "quality_flags": state.get("quality_flags", {}),
            "completed_at" : firestore.SERVER_TIMESTAMP,
        })

        # ── FCM push notification ──────────────────────────────────
        if job.get("fcm_token"):
            try:
                messaging.send(messaging.Message(
                    token        = job["fcm_token"],
                    notification = messaging.Notification(
                        title = "Report ready! 🎉",
                        body  = f"{job['title']} · Tap to download"
                    ),
                    data = {"job_id": jid, "download_url": download_url}
                ))
            except Exception:
                pass  # FCM failure should not fail the whole job

        # ── Email notification via Gmail SMTP ─────────────────────────
        notification_email = job.get("notification_email", "").strip()
        if notification_email:
            try:
                import os, smtplib
                from email.mime.multipart import MIMEMultipart
                from email.mime.text import MIMEText
                gmail_user = os.environ["GMAIL_USER"]
                gmail_pass = os.environ["GMAIL_APP_PASSWORD"]
                msg = MIMEMultipart("alternative")
                msg["Subject"] = f"Your report is ready: {job['title']}"
                msg["From"]    = f"ProjectDocs AI <{gmail_user}>"
                msg["To"]      = notification_email
                html = (
                    f"<h2 style='color:#6C63FF'>ProjectDocs AI — Report Ready</h2>"
                    f"<p>Hi {job.get('student_name', 'Student')},</p>"
                    f"<p>Your project report <strong>{job['title']}</strong> is ready.</p>"
                    f"<p><a href='{download_url}' style='background:#6C63FF;color:white;"
                    f"padding:12px 24px;border-radius:8px;text-decoration:none;display:inline-block'>"
                    f"Download Report (valid 24h)</a></p>"
                    f"<p style='color:#999;font-size:12px'>ProjectDocs AI · projdoc-aab8e.web.app</p>"
                )
                msg.attach(MIMEText(html, "html"))
                with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
                    smtp.login(gmail_user, gmail_pass)
                    smtp.sendmail(gmail_user, notification_email, msg.as_string())
                print(f"[EMAIL] Sent to {notification_email}")
            except Exception as email_err:
                print(f"[EMAIL] Failed: {email_err}")  # non-fatal

        return {**state, "drive_url": download_url}

    except Exception as e:
        if blob:
            try:
                blob.delete()
            except Exception:
                pass
        db.collection("jobs").document(jid).update({
            "status": "failed",
            "error" : str(e)
        })
        return {**state, "error": str(e)}


# ── Build the Graph ───────────────────────────────────────────────────────────

def build_report_graph():
    graph = StateGraph(ReportState)

    graph.add_node("planner",   planner_node)
    graph.add_node("generator", chapter_generator_node)
    graph.add_node("quality",   quality_node)
    graph.add_node("retry",     retry_node)
    graph.add_node("assembler", assembler_node)
    graph.add_node("appendix",  appendix_node)
    graph.add_node("delivery",  delivery_node)

    graph.set_entry_point("planner")
    graph.add_edge("planner",   "generator")
    graph.add_edge("generator", "quality")
    graph.add_conditional_edges(
        "quality",
        route_after_quality,
        {"retry": "retry", "assemble": "assembler"}
    )
    graph.add_edge("retry",     "quality")
    graph.add_edge("assembler", "appendix")
    graph.add_edge("appendix",  "delivery")
    graph.add_edge("delivery",  END)

    return graph.compile()


report_graph = build_report_graph()


def run_test_ping() -> dict:
    """Minimal LangGraph + Claude sanity check. Uses ~20 output tokens."""
    from typing import TypedDict as _TD
    from langgraph.graph import StateGraph as _SG, END as _END

    class _PS(_TD):
        response: str

    def _node(state: _PS) -> _PS:
        r = _get_client().messages.create(
            model      = "claude-haiku-4-5-20251001",
            max_tokens = 20,
            messages   = [{"role": "user", "content": "Reply with exactly: LangGraph+Claude OK"}],
        )
        return {"response": r.content[0].text.strip()}

    g = _SG(_PS)
    g.add_node("ping", _node)
    g.set_entry_point("ping")
    g.add_edge("ping", _END)
    out = g.compile().invoke({"response": ""})
    return {"status": "ok", "response": out["response"], "chars": len(out["response"])}


async def run_report_agent(job: dict) -> dict:
    """Called from main.py Cloud Function trigger."""
    initial_state: ReportState = {
        "job"             : job,
        "chapter_plan"    : {},
        "context_summary" : "",
        "chapters"        : {},
        "chapter_status"  : {},
        "retry_count"     : {},
        "quality_flags"   : {},
        "failed_chapters" : [],
        "docx_bytes"      : b"",
        "viva_content"    : "",
        "summary_content" : "",
        "drive_url"       : "",
        "error"           : "",
    }
    return await report_graph.ainvoke(initial_state)
